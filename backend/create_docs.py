import os
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
import docx
from docx.shared import Inches, Pt, RGBColor

def generate_pdf():
    pdf_filename = "MediKiosk_Bhashini_Project_Proposal.pdf"
    doc = SimpleDocTemplate(
        pdf_filename,
        pagesize=letter,
        rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54
    )
    
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontSize=22,
        leading=26,
        textColor=colors.HexColor('#0f766e'),
        spaceAfter=10
    )
    
    h2_style = ParagraphStyle(
        'DocH2',
        parent=styles['Heading2'],
        fontSize=14,
        leading=18,
        textColor=colors.HexColor('#0f172a'),
        spaceBefore=12,
        spaceAfter=6
    )
    
    body_style = ParagraphStyle(
        'DocBody',
        parent=styles['BodyText'],
        fontSize=10.5,
        leading=15,
        textColor=colors.HexColor('#334155'),
        spaceAfter=8
    )
    
    story = []
    
    story.append(Paragraph("MediKiosk — Multilingual AI Clinical Intake Platform", title_style))
    story.append(Paragraph("<b>Smart India Hackathon (SIH 2026) · Problem Statement 26047</b>", body_style))
    story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor('#0f766e'), spaceAfter=15))
    
    story.append(Paragraph("Executive Summary & Objective", h2_style))
    story.append(Paragraph(
        "MediKiosk is an AI-powered voice and touch pre-consultation patient intake platform designed specifically for high-volume Indian hospital Outpatient Departments (OPDs), which handle between 4,000 to 10,000 patients daily. Due to extreme doctor shortages, clinical consultation times are severely compressed to 2–5 minutes per patient, making complete clinical history taking difficult. MediKiosk bridges this gap by conducting adaptive pre-consultation clinical history taking in native regional Indian languages prior to the patient entering the consultation room.",
        body_style
    ))
    
    story.append(Paragraph("Integration of Bhashini AI Services", h2_style))
    story.append(Paragraph(
        "Bhashini Speech-to-Text (ASR) and Neural Machine Translation (NMT) serve as the core multilingual speech engine of MediKiosk. Patients interact with a Push-to-Talk (PTT) interface using native regional speech (Hindi, Tamil, Telugu, Marathi, Bengali, Gujarati, Kannada, etc.). Bhashini's ASR Dhruva API converts patient audio into native transcriptions, which are subsequently translated into standardized English clinical notes. Additionally, Bhashini Text-to-Speech (TTS) provides voice guidance for non-literate patients.",
        body_style
    ))
    
    story.append(Paragraph("Technical Solution Architecture", h2_style))
    story.append(Paragraph(
        "The software architecture consists of a React (Vite) Web Application utilizing the Web Audio API for Push-to-Talk audio recording, connected to a high-performance Python FastAPI backend. Audio payloads are processed asynchronously via HTTP REST calls to Bhashini's Dhruva pipeline. Extracted clinical symptoms and prior prescription OCR data are parsed into standardized FHIR-compliant JSON structures, generating a 30-second reviewable clinical summary for physicians.",
        body_style
    ))
    
    story.append(Paragraph("Compliance & Impact", h2_style))
    story.append(Paragraph(
        "MediKiosk is built in alignment with Ayushman Bharat Digital Mission (ABDM) ABHA sandbox standards and DPDP Act 2023 privacy regulations. Integrating Bhashini AI empowers patients across all linguistic backgrounds, reduces consultation overhead by up to 60%, and improves diagnostic accuracy across Indian healthcare institutions.",
        body_style
    ))
    
    doc.build(story)
    print(f"PDF saved: {os.path.abspath(pdf_filename)}")

def generate_docx():
    docx_filename = "MediKiosk_Bhashini_Project_Proposal.docx"
    doc = docx.Document()
    
    # Title
    p = doc.add_paragraph()
    run = p.add_run("MediKiosk — Multilingual AI Clinical Intake Platform\n")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 118, 110)
    
    run_sub = p.add_run("Smart India Hackathon (SIH 2026) · Problem Statement 26047\n")
    run_sub.font.size = Pt(12)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(71, 85, 105)
    
    sections = [
        ("Executive Summary & Objective", "MediKiosk is an AI-powered voice and touch pre-consultation patient intake platform designed specifically for high-volume Indian hospital Outpatient Departments (OPDs), which handle between 4,000 to 10,000 patients daily. Due to extreme doctor shortages, clinical consultation times are severely compressed to 2–5 minutes per patient, making complete clinical history taking difficult. MediKiosk bridges this gap by conducting adaptive pre-consultation clinical history taking in native regional Indian languages prior to the patient entering the consultation room."),
        ("Integration of Bhashini AI Services", "Bhashini Speech-to-Text (ASR) and Neural Machine Translation (NMT) serve as the core multilingual speech engine of MediKiosk. Patients interact with a Push-to-Talk (PTT) interface using native regional speech (Hindi, Tamil, Telugu, Marathi, Bengali, Gujarati, Kannada, etc.). Bhashini's ASR Dhruva API converts patient audio into native transcriptions, which are subsequently translated into standardized English clinical notes. Additionally, Bhashini Text-to-Speech (TTS) provides voice guidance for non-literate patients."),
        ("Technical Solution Architecture", "The software architecture consists of a React (Vite) Web Application utilizing the Web Audio API for Push-to-Talk audio recording, connected to a high-performance Python FastAPI backend. Audio payloads are processed asynchronously via HTTP REST calls to Bhashini's Dhruva pipeline. Extracted clinical symptoms and prior prescription OCR data are parsed into standardized FHIR-compliant JSON structures, generating a 30-second reviewable clinical summary for physicians."),
        ("Compliance & Impact", "MediKiosk is built in alignment with Ayushman Bharat Digital Mission (ABDM) ABHA sandbox standards and DPDP Act 2023 privacy regulations. Integrating Bhashini AI empowers patients across all linguistic backgrounds, reduces consultation overhead by up to 60%, and improves diagnostic accuracy across Indian healthcare institutions.")
    ]
    
    for title, content in sections:
        h = doc.add_paragraph()
        r = h.add_run(title)
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = RGBColor(15, 23, 42)
        
        cp = doc.add_paragraph()
        cr = cp.add_run(content)
        cr.font.size = Pt(11)
        cr.font.color.rgb = RGBColor(51, 65, 85)
        
    doc.save(docx_filename)
    print(f"DOCX saved: {os.path.abspath(docx_filename)}")

if __name__ == "__main__":
    generate_pdf()
    generate_docx()
